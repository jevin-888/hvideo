#!/usr/bin/env python3
# 编码声明：-*- coding: utf-8 -*-
"""
将Markdown文档转换为Word文档
需要安装: pip install python-docx markdown
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

def markdown_to_word(md_file, word_file):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 解析Markdown
    lines = md_content.split('\n')
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # 代码块处理
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_lines:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # 标题处理
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        # 表格分隔线
        elif line.startswith('---') or line.startswith('***'):
            continue
        # 空行
        elif not line.strip():
            doc.add_paragraph()
        # 列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        # 普通段落
        else:
            if line.strip():
                doc.add_paragraph(line)
    
    # 保存Word文档
    doc.save(word_file)
    print(f"Word文档已生成: {word_file}")

def main():
    docs_dir = Path(r'd:\Hvideo\docs')
    
    # 转换设计说明书
    design_md = docs_dir / '软件著作权-设计说明书.md'
    design_docx = docs_dir / '软件著作权-设计说明书.docx'
    if design_md.exists():
        markdown_to_word(design_md, design_docx)
    
    # 转换用户手册
    user_md = docs_dir / '软件著作权-用户手册.md'
    user_docx = docs_dir / '软件著作权-用户手册.docx'
    if user_md.exists():
        markdown_to_word(user_md, user_docx)
    
    # 源代码文档已经是txt格式，直接重命名为docx
    source_txt = docs_dir / '软件著作权-源代码文档.txt'
    source_docx = docs_dir / '软件著作权-源代码文档.docx'
    if source_txt.exists():
        # 创建Word文档并导入文本
        doc = Document()
        doc.styles['Normal'].font.name = 'Consolas'
        doc.styles['Normal'].font.size = Pt(9)
        
        with open(source_txt, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 按行添加
        for line in content.split('\n'):
            doc.add_paragraph(line)
        
        doc.save(source_docx)
        print(f"源代码文档已生成: {source_docx}")

if __name__ == '__main__':
    main()
